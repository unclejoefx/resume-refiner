"""Document parsing service with security hardening."""

import logging
import re
import regex  # For timeout support
from pathlib import Path
from typing import Optional, List, Dict, Any

import pdfplumber
from docx import Document
from email_validator import validate_email, EmailNotValidError

from app.models.resume import (
    ResumeContent,
    ContactInfo,
    Experience,
    Education,
    Skill
)
from app.services.parser_config import ParserConfig

# Configure logger
logger = logging.getLogger(__name__)


class ParsingError(Exception):
    """Custom exception for parsing errors."""
    pass


class DocumentParser:
    """Document parser for PDF and DOCX files with security hardening."""

    @staticmethod
    def parse(file_path: str, file_type: str) -> ResumeContent:
        """
        Parse document based on file type.

        Args:
            file_path: Path to document file
            file_type: File type (pdf or docx)

        Returns:
            ResumeContent object with extracted data

        Raises:
            ValueError: If file type is not supported
            ParsingError: If document cannot be parsed
        """
        logger.info(f"Starting parse for file: {file_path}, type: {file_type}")

        if file_type == "pdf":
            return DocumentParser.parse_pdf(file_path)
        elif file_type == "docx":
            return DocumentParser.parse_docx(file_path)
        else:
            error_msg = f"Unsupported file type: {file_type}"
            logger.error(error_msg)
            raise ValueError(error_msg)

    @staticmethod
    def parse_pdf(file_path: str) -> ResumeContent:
        """
        Parse PDF file and extract content.

        Args:
            file_path: Path to PDF file

        Returns:
            ResumeContent object with extracted data
        """
        try:
            raw_text = DocumentParser._extract_text_from_pdf(file_path)
            logger.info(f"Extracted {len(raw_text)} characters from PDF")
            return DocumentParser._parse_document_text(raw_text, file_path)

        except FileNotFoundError as e:
            logger.error(f"PDF file not found: {file_path}", exc_info=True)
            return ResumeContent(raw_text="File not found")

        except PermissionError as e:
            logger.error(f"Permission denied reading PDF: {file_path}", exc_info=True)
            return ResumeContent(raw_text="Permission denied")

        except Exception as e:
            logger.exception(f"Unexpected error parsing PDF: {file_path}")
            return ResumeContent(raw_text="Unable to process PDF file")

    @staticmethod
    def parse_docx(file_path: str) -> ResumeContent:
        """
        Parse DOCX file and extract content.

        Args:
            file_path: Path to DOCX file

        Returns:
            ResumeContent object with extracted data
        """
        try:
            raw_text = DocumentParser._extract_text_from_docx(file_path)
            logger.info(f"Extracted {len(raw_text)} characters from DOCX")
            return DocumentParser._parse_document_text(raw_text, file_path)

        except FileNotFoundError as e:
            logger.error(f"DOCX file not found: {file_path}", exc_info=True)
            return ResumeContent(raw_text="File not found")

        except PermissionError as e:
            logger.error(f"Permission denied reading DOCX: {file_path}", exc_info=True)
            return ResumeContent(raw_text="Permission denied")

        except Exception as e:
            logger.exception(f"Unexpected error parsing DOCX: {file_path}")
            return ResumeContent(raw_text="Unable to process DOCX file")

    @staticmethod
    def _extract_text_from_pdf(file_path: str) -> str:
        """Extract raw text from PDF file."""
        raw_text = ""
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    raw_text += page_text + "\n"

                # Prevent memory exhaustion
                if len(raw_text) > ParserConfig.MAX_RAW_TEXT_LENGTH:
                    logger.warning(f"PDF text exceeds max length, truncating at page {page_num}")
                    raw_text = raw_text[:ParserConfig.MAX_RAW_TEXT_LENGTH]
                    break

        return raw_text

    @staticmethod
    def _extract_text_from_docx(file_path: str) -> str:
        """Extract raw text from DOCX file."""
        doc = Document(file_path)
        raw_text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            raw_text += paragraph.text + "\n"

            # Prevent memory exhaustion
            if len(raw_text) > ParserConfig.MAX_RAW_TEXT_LENGTH:
                logger.warning("DOCX text exceeds max length, truncating")
                return raw_text[:ParserConfig.MAX_RAW_TEXT_LENGTH]

        # Extract text from tables (resumes often use tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    raw_text += cell.text + "\n"

                    # Prevent memory exhaustion
                    if len(raw_text) > ParserConfig.MAX_RAW_TEXT_LENGTH:
                        logger.warning("DOCX text exceeds max length, truncating")
                        return raw_text[:ParserConfig.MAX_RAW_TEXT_LENGTH]

        return raw_text

    @staticmethod
    def _parse_document_text(raw_text: str, file_path: str = "") -> ResumeContent:
        """
        Common parsing logic for all document types.

        Args:
            raw_text: Extracted text from document
            file_path: Optional file path for logging

        Returns:
            ResumeContent object with extracted data
        """
        if not raw_text.strip():
            logger.warning(f"Empty text extracted from {file_path}")
            return ResumeContent(raw_text="Unable to extract text from document")

        # Validate and sanitize text
        raw_text = DocumentParser._sanitize_text(raw_text)

        try:
            # Extract all components
            contact_info = DocumentParser._extract_contact_info(raw_text)
            summary = DocumentParser._extract_summary(raw_text)
            experience = DocumentParser._extract_experience(raw_text)
            education = DocumentParser._extract_education(raw_text)
            skills = DocumentParser._extract_skills(raw_text)
            sections = DocumentParser._identify_sections(raw_text)

            logger.info(f"Successfully parsed document: {len(experience)} experiences, "
                       f"{len(education)} education entries, {len(skills)} skill groups")

            return ResumeContent(
                contact_info=contact_info,
                summary=summary,
                experience=experience,
                education=education,
                skills=skills,
                raw_text=raw_text,
                sections=sections
            )

        except Exception as e:
            logger.exception(f"Error during document parsing for {file_path}")
            # Return partial results with error in sections
            return ResumeContent(
                raw_text=raw_text,
                sections={"error": "Partial parsing failure"}
            )

    @staticmethod
    def _sanitize_text(text: str) -> str:
        """Sanitize extracted text."""
        # Remove null bytes and control characters (except newlines and tabs)
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')

        # Normalize whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 consecutive newlines

        # Limit total length
        if len(text) > ParserConfig.MAX_RAW_TEXT_LENGTH:
            logger.warning(f"Text truncated from {len(text)} to {ParserConfig.MAX_RAW_TEXT_LENGTH}")
            text = text[:ParserConfig.MAX_RAW_TEXT_LENGTH]

        return text

    @staticmethod
    def _extract_contact_info(text: str) -> Optional[ContactInfo]:
        """Extract and validate contact information from text."""
        # Use only first portion for contact info (usually at top)
        search_text = text[:ParserConfig.MAX_TEXT_FOR_CONTACT_EXTRACTION]

        contact = ContactInfo()

        try:
            # Extract and validate email
            email_match = re.search(ParserConfig.EMAIL_PATTERN, search_text)
            if email_match:
                try:
                    valid = validate_email(email_match.group(0), check_deliverability=False)
                    contact.email = valid.normalized
                    logger.debug(f"Extracted email: {contact.email}")
                except EmailNotValidError as e:
                    logger.debug(f"Invalid email found: {email_match.group(0)}")

            # Extract phone number
            phone_match = re.search(ParserConfig.PHONE_PATTERN, search_text)
            if phone_match:
                contact.phone = DocumentParser._normalize_phone(phone_match.group(0))
                logger.debug(f"Extracted phone: {contact.phone}")

            # Extract LinkedIn URL
            linkedin_match = re.search(ParserConfig.LINKEDIN_PATTERN, search_text, re.IGNORECASE)
            if linkedin_match:
                linkedin_url = linkedin_match.group(0)
                if not linkedin_url.startswith('http'):
                    linkedin_url = 'https://' + linkedin_url
                contact.linkedin = linkedin_url
                logger.debug(f"Extracted LinkedIn: {contact.linkedin}")

            # Extract name (heuristic: first non-empty line that's not too long)
            lines = [line.strip() for line in search_text.split('\n') if line.strip()]
            for line in lines[:5]:  # Check first 5 lines
                # Skip lines that look like section headers or emails
                if (len(line) < 50 and
                    not line.isupper() and
                    '@' not in line and
                    not any(char.isdigit() for char in line)):
                    contact.name = line
                    logger.debug(f"Extracted name: {contact.name}")
                    break

        except Exception as e:
            logger.warning(f"Error extracting contact info: {e}")

        return contact if any([contact.email, contact.phone, contact.name]) else None

    @staticmethod
    def _normalize_phone(phone: str) -> str:
        """Normalize phone number format."""
        # Remove all non-digit characters except +
        digits = ''.join(c for c in phone if c.isdigit() or c == '+')
        return digits

    @staticmethod
    def _extract_summary(text: str) -> Optional[str]:
        """Extract professional summary from text."""
        try:
            for header in ParserConfig.SECTION_HEADERS['summary']:
                # Simplified pattern to reduce ReDoS risk
                pattern = rf'(?:^|\n)({re.escape(header)})\s*[:\-]?\s*\n'

                try:
                    # Use regex library with timeout
                    match = regex.search(
                        pattern,
                        text,
                        flags=regex.IGNORECASE,
                        timeout=ParserConfig.REGEX_TIMEOUT
                    )

                    if match:
                        # Extract text after header until next section
                        start = match.end()
                        # Look for next section header
                        next_section = len(text)
                        for section_headers in ParserConfig.SECTION_HEADERS.values():
                            for next_header in section_headers:
                                next_pattern = rf'\n{re.escape(next_header)}\s*[:\-]?\s*\n'
                                try:
                                    next_match = regex.search(
                                        next_pattern,
                                        text[start:start+2000],  # Limit search
                                        flags=regex.IGNORECASE,
                                        timeout=ParserConfig.REGEX_TIMEOUT
                                    )
                                    if next_match:
                                        potential_end = start + next_match.start()
                                        next_section = min(next_section, potential_end)
                                except TimeoutError:
                                    logger.warning("Regex timeout finding next section")

                        summary_text = text[start:next_section].strip()

                        # Clean up and validate
                        summary_text = ' '.join(summary_text.split())

                        if len(summary_text) >= ParserConfig.MIN_SUMMARY_LENGTH:
                            summary = summary_text[:ParserConfig.MAX_SUMMARY_LENGTH]
                            logger.debug(f"Extracted summary: {len(summary)} chars")
                            return summary

                except TimeoutError:
                    logger.warning(f"Regex timeout on summary pattern for header: {header}")
                    continue

        except Exception as e:
            logger.warning(f"Error extracting summary: {e}")

        return None

    @staticmethod
    def _extract_experience(text: str) -> List[Experience]:
        """Extract work experience from text."""
        experiences = []

        try:
            for header in ParserConfig.SECTION_HEADERS['experience']:
                pattern = rf'(?:^|\n)({re.escape(header)})\s*[:\-]?\s*\n'

                try:
                    match = regex.search(
                        pattern,
                        text,
                        flags=regex.IGNORECASE,
                        timeout=ParserConfig.REGEX_TIMEOUT
                    )

                    if match:
                        start = match.end()

                        # Find end of experience section
                        end = len(text)
                        for next_header in (ParserConfig.SECTION_HEADERS['education'] +
                                          ParserConfig.SECTION_HEADERS['skills']):
                            next_pattern = rf'\n{re.escape(next_header)}\s*[:\-]?\s*\n'
                            try:
                                next_match = regex.search(
                                    next_pattern,
                                    text[start:start+5000],
                                    flags=regex.IGNORECASE,
                                    timeout=ParserConfig.REGEX_TIMEOUT
                                )
                                if next_match:
                                    potential_end = start + next_match.start()
                                    end = min(end, potential_end)
                            except TimeoutError:
                                pass

                        experience_text = text[start:end].strip()

                        if experience_text:
                            # For now, create single entry with full text
                            # TODO: Enhance to split into individual job entries
                            description = experience_text[:ParserConfig.MAX_EXPERIENCE_DESCRIPTION_LENGTH]

                            exp = Experience(
                                company="Multiple positions listed",
                                position="See description",
                                description=description,
                                bullets=[]
                            )
                            experiences.append(exp)
                            logger.debug("Extracted experience section")
                        break

                except TimeoutError:
                    logger.warning(f"Regex timeout extracting experience for header: {header}")

        except Exception as e:
            logger.warning(f"Error extracting experience: {e}")

        return experiences[:ParserConfig.MAX_EXPERIENCE_ENTRIES]

    @staticmethod
    def _extract_education(text: str) -> List[Education]:
        """Extract education from text."""
        education_list = []

        try:
            for header in ParserConfig.SECTION_HEADERS['education']:
                pattern = rf'(?:^|\n)({re.escape(header)})\s*[:\-]?\s*\n'

                try:
                    match = regex.search(
                        pattern,
                        text,
                        flags=regex.IGNORECASE,
                        timeout=ParserConfig.REGEX_TIMEOUT
                    )

                    if match:
                        start = match.end()

                        # Find end of education section
                        end = len(text)
                        for next_header in ParserConfig.SECTION_HEADERS['skills']:
                            next_pattern = rf'\n{re.escape(next_header)}\s*[:\-]?\s*\n'
                            try:
                                next_match = regex.search(
                                    next_pattern,
                                    text[start:start+3000],
                                    flags=regex.IGNORECASE,
                                    timeout=ParserConfig.REGEX_TIMEOUT
                                )
                                if next_match:
                                    end = start + next_match.start()
                            except TimeoutError:
                                pass

                        education_text = text[start:end].strip()

                        # Look for degree patterns
                        for pattern in ParserConfig.DEGREE_PATTERNS:
                            try:
                                matches = regex.finditer(
                                    pattern,
                                    education_text,
                                    flags=regex.IGNORECASE,
                                    timeout=ParserConfig.REGEX_TIMEOUT
                                )
                                for match in matches:
                                    institution_text = match.group(0).strip()
                                    institution_text = institution_text[:ParserConfig.MAX_EDUCATION_INSTITUTION_LENGTH]

                                    edu = Education(
                                        institution=institution_text,
                                        degree="",
                                        achievements=[]
                                    )
                                    education_list.append(edu)

                                    if len(education_list) >= ParserConfig.MAX_EDUCATION_ENTRIES:
                                        break
                            except TimeoutError:
                                logger.warning("Regex timeout in education pattern matching")

                        # If no matches, create generic entry
                        if not education_list and education_text:
                            edu = Education(
                                institution=education_text[:ParserConfig.MAX_EDUCATION_INSTITUTION_LENGTH],
                                degree="",
                                achievements=[]
                            )
                            education_list.append(edu)

                        logger.debug(f"Extracted {len(education_list)} education entries")
                        break

                except TimeoutError:
                    logger.warning(f"Regex timeout extracting education for header: {header}")

        except Exception as e:
            logger.warning(f"Error extracting education: {e}")

        return education_list[:ParserConfig.MAX_EDUCATION_ENTRIES]

    @staticmethod
    def _extract_skills(text: str) -> List[Skill]:
        """Extract skills from text."""
        skills_list = []

        try:
            for header in ParserConfig.SECTION_HEADERS['skills']:
                pattern = rf'(?:^|\n)({re.escape(header)})\s*[:\-]?\s*\n'

                try:
                    match = regex.search(
                        pattern,
                        text,
                        flags=regex.IGNORECASE,
                        timeout=ParserConfig.REGEX_TIMEOUT
                    )

                    if match:
                        start = match.end()
                        # Limit skills section search
                        skills_text = text[start:start+1000].strip()

                        # Split by common delimiters
                        skill_items = re.split(r'[,;•·\n]+', skills_text)
                        skill_items = [
                            s.strip() for s in skill_items
                            if s.strip() and len(s.strip()) >= ParserConfig.MIN_SKILL_LENGTH
                        ]

                        if skill_items:
                            skill = Skill(
                                category="Technical Skills",
                                skills=skill_items[:ParserConfig.MAX_SKILLS_COUNT]
                            )
                            skills_list.append(skill)
                            logger.debug(f"Extracted {len(skill_items)} skills")
                        break

                except TimeoutError:
                    logger.warning(f"Regex timeout extracting skills for header: {header}")

        except Exception as e:
            logger.warning(f"Error extracting skills: {e}")

        return skills_list

    @staticmethod
    def _identify_sections(text: str) -> Dict[str, Any]:
        """Identify all sections in the document."""
        sections = {}

        try:
            for section_type, headers in ParserConfig.SECTION_HEADERS.items():
                for header in headers:
                    pattern = rf'(?:^|\n)({re.escape(header)})\s*[:\-]?'

                    try:
                        match = regex.search(
                            pattern,
                            text,
                            flags=regex.IGNORECASE,
                            timeout=ParserConfig.REGEX_TIMEOUT
                        )

                        if match:
                            sections[section_type] = {
                                'found': True,
                                'position': match.start(),
                                'header': match.group(1)
                            }
                            break

                    except TimeoutError:
                        logger.warning(f"Regex timeout identifying section: {header}")
                        continue

        except Exception as e:
            logger.warning(f"Error identifying sections: {e}")

        return sections
